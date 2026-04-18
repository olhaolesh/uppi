"""Baseline-тести для генератора атестації та DOCX template filler."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from itemadapter import ItemAdapter

import uppi.docs.attestazione_template_filler as legacy_template_filler
import uppi.services.attestazione_template_filler as canonical_template_filler
from uppi.docs.attestazione_template_filler import (
    fill_attestazione_template,
    fill_underscored,
    underscored,
)
from uppi.domain.immobile import Immobile
from uppi.services.attestazione_generator import build_template_params


def _make_immobile(**overrides) -> Immobile:
    """Створює тестові дані для сценаріїв цього модуля."""
    base = {
        "foglio": "12",
        "numero": "345",
        "sub": "7",
        "rendita": "€ 123.45",
        "superficie_totale": 98.7,
        "categoria": "A/2",
    }
    base.update(overrides)
    return Immobile(**base)


def _make_adapter(**overrides) -> ItemAdapter:
    """Створює тестові дані для сценаріїв цього модуля."""
    base = {
        "locatore_cf": " RSSMRA80A01H501Z ",
        "contratto_data": "01/01/2025",
        "decorrenza_data": "01/02/2025",
        "registrazione_data": "05/02/2025",
        "registrazione_num": "123",
        "agenzia_entrate_sede": "PESCARA",
        "conduttore_nome": "Mario Bianchi",
        "conduttore_cf": "BNCMRA80A01H501Z",
        "conduttore_comune": "Pescara",
        "conduttore_via": "Via Test 12",
        "canone_contrattuale_mensile": "650",
    }
    base.update(overrides)
    return ItemAdapter(base)


def _make_contract_ctx(**overrides) -> dict:
    """Створює тестові дані для сценаріїв цього модуля."""
    base = {
        "overrides": {
            "locatore_comune_res": "Chieti",
            "immobile_via_override": "Corso Roma",
        },
        "elements": {
            "a1": "X",
            "b2": "X",
            "d5": "Y",
        },
        "contract": {
            "contratto_data": "DB CONTRATTO",
            "registrazione_num": "DB-REG",
            "canone_contrattuale_mensile": 700,
        },
        "parties": {
            "LOCATORE": {
                "cf": "RSSMRA80A01H501Z",
                "name": "MARIO",
                "surname": "ROSSI",
                "address": {
                    "comune": "Pescara",
                    "via_full": "Via Roma",
                    "civico": "10",
                },
            },
            "CONDUTTORE": {
                "name": "DB Conduttore",
                "cf": "DBCF123",
            },
        },
        "immobile": {
            "comune": "Pescara",
            "via": "Via Milano",
            "civico": "12",
            "piano": "3",
            "interno": "4",
        },
    }
    base.update(overrides)
    return base


def _make_minimal_template(template_path: Path) -> None:
    """Створює тестові дані для сценаріїв цього модуля."""
    doc = Document()

    p1 = doc.add_paragraph()
    p1.add_run("Locatore: ")
    p1.add_run("{{LOCATORE_NOME}}")

    p2 = doc.add_paragraph()
    p2.add_run("Comune residenza: ")
    p2.add_run("{{LOCATORE_COMUNE_RES}}")

    p3 = doc.add_paragraph()
    p3.add_run("Conduttore CF: ")
    p3.add_run("{{CONDUTTORE_CF}}")

    p4 = doc.add_paragraph()
    p4.add_run("Decorrenza: ")
    p4.add_run("{{DECORRENZA_DATA}}")

    p5 = doc.add_paragraph()
    p5.add_run("Zona: ")
    p5.add_run("{{CAN_ZONA}}")

    p6 = doc.add_paragraph()
    p6.add_run("Split: ")
    p6.add_run("{{LOC")
    p6.add_run("ATORE_NOME}}")

    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Immobile via: {{IMMOBILE_VIA}}"

    doc.save(template_path)


def _paragraph_texts(doc: Document) -> list[str]:
    """Допоміжний тестовий хелпер для цього модуля."""
    return [paragraph.text for paragraph in doc.paragraphs]


def test_attestazione_generator_happy_path_populates_current_core_params():
    """Перевіряє сценарій, описаний у назві тесту."""
    params = build_template_params(
        _make_adapter(),
        _make_immobile(),
        _make_contract_ctx(),
    )

    assert params["{{LOCATORE_CF}}"] == "RSSMRA80A01H501Z"
    assert params["{{LOCATORE_NOME}}"] == "Mario Rossi"
    assert params["{{LOCATORE_COMUNE_RES}}"] == "Chieti"
    assert params["{{LOCATORE_VIA}}"] == "Via Roma"
    assert params["{{IMMOBILE_COMUNE}}"] == "Pescara"
    assert params["{{IMMOBILE_VIA}}"] == "Corso Roma"
    assert params["{{IMMOBILE_CIVICO}}"] == "12"
    assert params["{{IMMOBILE_PIANO}}"] == "3"
    assert params["{{IMMOBILE_INTERNO}}"] == "4"
    assert params["{{FOGLIO}}"] == "12"
    assert params["{{NUMERO}}"] == "345"
    assert params["{{SUB}}"] == "7"
    assert params["{{RENDITA}}"] == "€ 123.45"
    assert params["{{SUPERFICIE_TOTALE}}"] == "98.7"
    assert params["{{APP_FOGL}}"] == "12"
    assert params["{{APP_PART}}"] == "345"
    assert params["{{APP_SUB}}"] == "7"
    assert params["{{APP_REND}}"] == "€ 123.45"
    assert params["{{APP_SCAT}}"] == "98.7"
    assert params["{{APP_SRIP}}"] == "X"
    assert params["{{APP_CAT}}"] == "A/2"
    assert params["{{GAR_FOGL}}"] == ""
    assert params["{{PST_FOGL}}"] == ""
    assert params["{{TOT_SCAT}}"] == "98.7"
    assert params["{{TOT_SRIP}}"] == "X"
    assert params["{{TOT_CAT}}"] == "X"
    assert params["{{A1}}"] == "X"
    assert params["{{a1}}"] == "X"
    assert params["{{B2}}"] == "X"
    assert params["{{D5}}"] == "Y"
    assert params["{{A_CNT}}"] == "1"
    assert params["{{B_CNT}}"] == "1"
    assert params["{{C_CNT}}"] == "0"
    assert params["{{D_CNT}}"] == "1"
    assert params["{{CAN_MENSILE}}"] == "650.00"


def test_attestazione_template_filler_compatibility_shim_keeps_old_and_new_import_paths_working():
    """Перевіряє, що migration shim не ламає старий import path."""
    assert legacy_template_filler.fill_attestazione_template is canonical_template_filler.fill_attestazione_template
    assert legacy_template_filler.fill_underscored is canonical_template_filler.fill_underscored
    assert legacy_template_filler.underscored is canonical_template_filler.underscored


def test_attestazione_generator_known_current_behavior_yaml_fields_override_db_contract_and_conduttore_state():
    """Перевіряє сценарій, описаний у назві тесту."""
    adapter = _make_adapter(
        conduttore_nome="YAML Conduttore",
        conduttore_cf="YAMLCF123",
        registrazione_num="YAML-REG",
    )
    contract_ctx = _make_contract_ctx(
        contract={
            "registrazione_num": "DB-REG",
            "canone_contrattuale_mensile": 700,
        },
        parties={
            "LOCATORE": {
                "cf": "RSSMRA80A01H501Z",
                "name": "MARIO",
                "surname": "ROSSI",
                "address": {
                    "comune": "Pescara",
                    "via_full": "Via Roma",
                    "civico": "10",
                },
            },
            "CONDUTTORE": {
                "name": "DB Conduttore",
                "cf": "DBCF123",
            },
        },
    )

    params = build_template_params(adapter, _make_immobile(), contract_ctx)

    assert params["{{CONDUTTORE_NOME}}"] == "YAML Conduttore"
    assert params["{{CONDUTTORE_CF}}"] == "YAMLCF123"
    assert params["{{REGISTRAZIONE_NUM}}"] == "YAML-REG"
    assert params["{{CAN_MENSILE}}"] == "650.00"


def test_attestazione_generator_partial_input_returns_current_blank_placeholders():
    """Перевіряє сценарій, описаний у назві тесту."""
    params = build_template_params(
        ItemAdapter({"locatore_cf": "RSSMRA80A01H501Z"}),
        Immobile(),
        {"elements": {}, "parties": {}, "contract": {}, "overrides": {}, "immobile": {}},
    )

    assert params["{{LOCATORE_CF}}"] == "RSSMRA80A01H501Z"
    assert params["{{LOCATORE_NOME}}"] == ""
    assert params["{{LOCATORE_COMUNE_RES}}"] == ""
    assert params["{{IMMOBILE_VIA}}"] == ""
    assert params["{{APP_FOGL}}"] == ""
    assert params["{{APP_SCAT}}"] == ""
    assert params["{{CONDUTTORE_NOME}}"] == ""
    assert params["{{REGISTRAZIONE_NUM}}"] == ""
    assert params["{{A_CNT}}"] == "0"
    assert params["{{B_CNT}}"] == "0"
    assert params["{{C_CNT}}"] == "0"
    assert params["{{D_CNT}}"] == "0"
    assert params["{{CAN_ZONA}}"] == ""
    assert params["{{CAN_MENSILE}}"] == ""


def test_attestazione_generator_does_not_use_db_canone_as_run_only_default():
    """Blank run-only canone must stay blank instead of falling back to stored contract data."""
    params = build_template_params(
        ItemAdapter({"locatore_cf": "RSSMRA80A01H501Z"}),
        _make_immobile(),
        _make_contract_ctx(contract={"canone_contrattuale_mensile": 700}),
    )

    assert params["{{CAN_MENSILE}}"] == ""


def test_attestazione_generator_treats_run_only_clear_markers_as_blank_current_run_state():
    """Run-only `-` must not leak into DOCX placeholders as a literal dash."""
    params = build_template_params(
        _make_adapter(
            conduttore_cf="-",
            decorrenza_data="-",
            registrazione_num="-",
            canone_contrattuale_mensile="-",
        ),
        _make_immobile(),
        _make_contract_ctx(),
    )

    assert params["{{CONDUTTORE_CF}}"] == ""
    assert params["{{DECORRENZA_DATA}}"] == ""
    assert params["{{REGISTRAZIONE_NUM}}"] == ""
    assert params["{{CAN_MENSILE}}"] == ""


def test_attestazione_generator_known_current_behavior_ignore_surcharges_keeps_transitorio_text_but_excludes_it_from_var_range():
    """Перевіряє сценарій, описаний у назві тесту."""
    contract_ctx = {
        "canone_calc": {
            "canone_input": {
                "istat": 5,
                "ignore_surcharges": True,
                "superficie_catastale": 80,
                "arredato": 0.10,
                "energy_class": "b",
                "durata_anni": 4,
                "contract_kind": "domain.TRANSITORIO",
            },
            "result": {
                "zona": "B1",
                "subfascia": "ALTA",
                "base_euro_mq_istat": 12.5,
                "canone_base_annuo": 1000,
                "base_min_euro_mq": 10,
                "base_max_euro_mq": 20,
            },
        }
    }

    params = build_template_params(ItemAdapter({}), _make_immobile(superficie_totale=100), contract_ctx)

    assert params["{{CAN_ISTAT}}"] == "ISTAT (+5.00%)"
    assert params["{{CAN_ZONA}}"] == "B1"
    assert params["{{CAN_SUBFASCIA}}"] == "ALTA"
    assert params["{{CAN_MQ}}"] == "80"
    assert params["{{CAN_MQ_ANNUO}}"] == "12.50"
    assert params["{{CAN_TOTALE_ANNUO}}"] == "1000.00"
    assert params["{{CAN_ARREDATO}}"] == "(+10%)  +100.00"
    assert params["{{CAN_CLASSE_B}}"] == "(+4%) +40.00"
    assert params["{{CAN_DURATA}}"] == "(+5%) +50.00"
    assert params["{{CAN_TRANSITORIO}}"] == "(+15%) +150.00"
    assert params["{{CAN_ANNUO_VAR_MIN}}"] == "952.00"
    assert params["{{CAN_ANNUO_VAR_MAX}}"] == "1904.00"
    assert params["{{CAN_MENSILE_VAR_MIN}}"] == "79.33"
    assert params["{{CAN_MENSILE_VAR_MAX}}"] == "158.67"


def test_template_filler_happy_path_creates_current_docx_output(tmp_path):
    """Перевіряє сценарій, описаний у назві тесту."""
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "out"
    _make_minimal_template(template_path)

    params = build_template_params(
        _make_adapter(),
        _make_immobile(),
        _make_contract_ctx(),
    )
    params["{{CAN_ZONA}}"] = "B1"

    output_path = fill_attestazione_template(
        str(template_path),
        str(output_dir),
        "attestazione.docx",
        params,
        underscored,
    )

    assert output_path == str(output_dir / "attestazione.docx")
    assert Path(output_path).exists()

    document = Document(output_path)
    paragraph_texts = _paragraph_texts(document)
    table_text = document.tables[0].cell(0, 0).text

    assert "Mario Rossi" in paragraph_texts[0]
    assert "Chieti" in paragraph_texts[1]
    assert "BNCMRA80A01H501Z" in paragraph_texts[2]
    assert "01/02/2025" in paragraph_texts[3]
    assert paragraph_texts[4] == "Zona: B1"
    assert table_text == "Immobile via: " + fill_underscored("Corso Roma", underscored["{{IMMOBILE_VIA}}"])


def test_template_filler_partial_input_uses_current_underscores_and_blank_removal(tmp_path):
    """Перевіряє сценарій, описаний у назві тесту."""
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "out"
    _make_minimal_template(template_path)

    output_path = fill_attestazione_template(
        str(template_path),
        str(output_dir),
        "attestazione_partial.docx",
        {
            "{{LOCATORE_NOME}}": "Mario Rossi",
        },
        underscored,
    )

    document = Document(output_path)
    paragraph_texts = _paragraph_texts(document)
    table_text = document.tables[0].cell(0, 0).text

    assert "Mario Rossi" in paragraph_texts[0]
    assert paragraph_texts[1] == "Comune residenza: " + ("_" * underscored["{{LOCATORE_COMUNE_RES}}"])
    assert paragraph_texts[2] == "Conduttore CF: " + ("_" * underscored["{{CONDUTTORE_CF}}"])
    assert paragraph_texts[3] == "Decorrenza: " + ("_" * underscored["{{DECORRENZA_DATA}}"])
    assert paragraph_texts[4] == "Zona: "
    assert table_text == "Immobile via: " + ("_" * underscored["{{IMMOBILE_VIA}}"])


def test_template_filler_known_current_behavior_placeholder_split_across_runs_is_not_replaced(tmp_path):
    """Перевіряє сценарій, описаний у назві тесту."""
    template_path = tmp_path / "template.docx"
    output_dir = tmp_path / "out"
    _make_minimal_template(template_path)

    output_path = fill_attestazione_template(
        str(template_path),
        str(output_dir),
        "attestazione_split.docx",
        {
            "{{LOCATORE_NOME}}": "Mario Rossi",
        },
        underscored,
    )

    document = Document(output_path)

    assert document.paragraphs[5].text == "Split: {{LOCATORE_NOME}}"


def test_fill_underscored_known_current_behavior_long_text_is_returned_without_truncation():
    """Перевіряє сценарій, описаний у назві тесту."""
    long_value = "VALORE-MOLTO-LUNGO"

    assert fill_underscored(long_value, 4) == long_value
