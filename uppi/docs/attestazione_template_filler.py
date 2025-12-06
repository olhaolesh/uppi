import shutil
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_UNDERLINE

# Співвідношення ширини пробілу до ширини "_"
SPACE_TO_UNDERSCORE_RATIO = 1.8

# Регулярка для пошуку всіх плейсхолдерів типу {{KEY}}
PLACEHOLDER_RE = re.compile(r"{{[^}]+}}")


def fill_underscored(text: str | None, length: int) -> str:
    """
    Повертає рядок, який зберігає "ширину" вихідних підкреслень у шаблоні.

    - Якщо text = None або порожній → просто повертаємо "_" * length
      (класичний "порожній бланк").
    - Якщо text є → вставляємо його всередину, а навколо додаємо пробіли,
      щоб в сумі приблизно співпало з length підкресленнями.
    """
    # 1. Якщо даних немає — повертаємо оригінальні підкреслення
    if not text:
        return "_" * max(length, 0)

    text = str(text).strip()

    # Кількість символів, яку треба заповнити в оригінальній метриці '_'
    original_padding_length = length - len(text)

    # Якщо текст довший за "бланк" — просто повертаємо текст як є
    if original_padding_length <= 0:
        return text

    # Скільки пробілів треба для збереження ширини (грубо)
    target_space_padding = int(original_padding_length * SPACE_TO_UNDERSCORE_RATIO)

    # Центрування
    left_padding = target_space_padding // 2
    right_padding = target_space_padding - left_padding

    return " " * left_padding + text + " " * right_padding


def _replace_in_run(run, params: dict, underscored: dict):
    """
    Обробляє ОДИН run:

    - знаходить усі {{KEY}} у run.text;
    - для кожного KEY:
        * якщо KEY == {{CONDUTTORE_CF}}:
              спеціальна логіка — просто вставляємо значення,
              щоб не ламати верстку, underline лишається з шаблону.
        * якщо KEY у params:
              - якщо KEY у underscored → fill_underscored(value, length)
              - інакше → str(value або "")
        * якщо KEY НЕ у params:
              - якщо KEY у underscored → "порожній бланк": "_" * length
              - інакше → "" (прибрати повністю).
    - якщо хоч один KEY був з underscored → ставимо underline на run.
    """
    text = run.text
    if "{{" not in text:
        return  # нічого робити

    underline_needed = False

    def repl(match: re.Match) -> str:
        nonlocal underline_needed
        key = match.group(0)  # включно з {{ }}

        # Значення з params (може бути None / "")
        value = params.get(key, None)

        # --- СПЕЦІАЛЬНИЙ ВИПАДОК ДЛЯ {{CONDUTTORE_CF}} ---
        if key == "{{CONDUTTORE_CF}}":
            underline_needed = True
            length = underscored.get(key, 0)
            # якщо є значення → просто вставляємо текст
            if value:
                return str(value)
            # якщо немає значення → малюємо рівно length підкреслень
            return "_" * max(length, 0)

        # Плейсхолдер з фіксованою довжиною підкреслень
        if key in underscored:
            underline_needed = True
            length = underscored[key]
            return fill_underscored(value, length)

        # Звичайний плейсхолдер, без підкреслень
        if value is not None:
            return str(value)

        # Немає в params і не underscored → прибираємо ключ
        return ""

    new_text = PLACEHOLDER_RE.sub(repl, text)
    if new_text != text:
        run.text = new_text
        if underline_needed:
            run.font.underline = WD_UNDERLINE.SINGLE


def replace_in_paragraph(paragraph, params: dict, underscored: dict):
    """Обробляє всі run-и в параграфі."""
    for run in paragraph.runs:
        _replace_in_run(run, params, underscored)


def replace_in_cell(cell, params: dict, underscored: dict):
    """Обробляє всі параграфи в комірці таблиці."""
    for paragraph in cell.paragraphs:
        replace_in_paragraph(paragraph, params, underscored)


def fill_attestazione_template(
    template_path: str,
    output_folder: str,
    filename: str,
    params: dict,
    underscored: dict,
) -> str:
    """
    Основна функція:

    - робить копію шаблону в output_folder/filename;
    - проганяє всі параграфи та таблиці, замінюючи {{KEY}} згідно params/underscored;
    - повертає шлях до заповненого файлу.
    """
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    out_path = output_folder / filename

    # 1. копія шаблону
    shutil.copy(template_path, out_path)

    # 2. завантажуємо документ
    doc = Document(out_path)

    # 3. параграфи
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, params, underscored)

    # 4. таблиці
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, params, underscored)

    # 5. зберігаємо
    doc.save(out_path)
    return str(out_path)


# ----------------------------------------------
# Фіксована довжина підкреслень для відповідних ключів
# ----------------------------------------------
underscored = {
    "{{LOCATORE_NOME}}": 40,
    "{{LOCATORE_CF}}": 25,
    "{{LOCATORE_COMUNE_RES}}": 27,
    "{{LOCATORE_VIA}}": 27,
    "{{LOCATORE_CIVICO}}": 4,
    "{{IMMOBILE_COMUNE}}": 24,
    "{{IMMOBILE_VIA}}": 27,
    "{{IMMOBILE_CIVICO}}": 4,
    "{{IMMOBILE_PIANO}}": 4,
    "{{IMMOBILE_INTERNO}}": 4,
    "{{CONTRATTO_DATA}}": 13,
    "{{CONDUTTORE_NOME}}": 27,
    "{{CONDUTTORE_CF}}": 21,
    "{{CONDUTTORE_COMUNE}}": 24,
    "{{CONDUTTORE_VIA}}": 27,
    "{{DECORRENZA_DATA}}": 18,
    "{{REGISTRAZIONE_DATA}}": 13,
    "{{REGISTRAZIONE_NUM}}": 4,
    "{{AGENZIA_ENTRATE_SEDE}}": 25,
}


if __name__ == "__main__":
    # Основний блок для тестування
    BASE_DIR = Path(__file__).resolve().parent.parent.parent
    template_path = BASE_DIR / "attestazione_template" / "template_attestazione_pescara.docx"

    output_dir = "output"
    filename = "attestazione_filled.docx"

    params = {
        "{{LOCATORE_NOME}}": "Mario Rossi",
        "{{LOCATORE_CF}}": "RSSMRA80A01H501X",
        "{{LOCATORE_COMUNE_RES}}": "Pescara",
        "{{LOCATORE_VIA}}": "Predazzo",
        "{{LOCATORE_CIVICO}}": 43,
        "{{IMMOBILE_COMUNE}}": "Montesilvano",
        "{{IMMOBILE_VIA}}": "C-so Umberto I",
        "{{IMMOBILE_CIVICO}}": "316",
        "{{IMMOBILE_PIANO}}": "4",
        "{{IMMOBILE_INTERNO}}": "",
        "{{CONTRATTO_DATA}}": "",
        "{{CONDUTTORE_NOME}}": "Biaocchi Giovana",
        "{{CONDUTTORE_CF}}": "BCCGNN44M45G488W",
        "{{CONDUTTORE_COMUNE}}": "",
        "{{CONDUTTORE_VIA}}": "",
        "{{DECORRENZA_DATA}}": "15/10/2025",
        "{{REGISTRAZIONE_DATA}}": "",
        "{{REGISTRAZIONE_NUM}}": "",
        "{{AGENZIA_ENTRATE_SEDE}}": "",
    }

    print(BASE_DIR)
    filled_path = fill_attestazione_template(
        template_path,
        output_dir,
        filename,
        params,
        underscored,
    )
    print(f"Filled document saved to: {filled_path}")