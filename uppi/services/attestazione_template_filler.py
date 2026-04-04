"""Canonical runtime-home для DOCX template filler атестації."""

from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_UNDERLINE

# Співвідношення ширини пробілу до ширини "_"
SPACE_TO_UNDERSCORE_RATIO = 1.8

# Регулярка для пошуку всіх плейсхолдерів типу {{KEY}}
PLACEHOLDER_RE = re.compile(r"{{[^}]+}}")


def fill_underscored(text: str | None, length: int) -> str:
    """
    Повертає рядок, який зберігає "ширину" вихідних підкреслень у шаблоні.

    - Якщо text = None або порожній -> просто повертаємо "_" * length
      (класичний "порожній бланк").
    - Якщо text є -> вставляємо його всередину, а навколо додаємо пробіли,
      щоб в сумі приблизно співпало з length підкресленнями.
    """
    if not text:
        return "_" * max(length, 0)

    text = str(text).strip()
    original_padding_length = length - len(text)

    if original_padding_length <= 0:
        return text

    target_space_padding = int(original_padding_length * SPACE_TO_UNDERSCORE_RATIO)
    left_padding = target_space_padding // 2
    right_padding = target_space_padding - left_padding

    return " " * left_padding + text + " " * right_padding


def _replace_in_run(run, params: dict, underscored: dict):
    """
    Обробляє один run і замінює плейсхолдери згідно current DOCX contract.

    Цей модуль є migration target для production code з `uppi/docs/`.
    У межах compatibility-slice semantics заміни плейсхолдерів не змінюються.
    """
    text = run.text
    if "{{" not in text:
        return

    underline_needed = False

    def repl(match: re.Match) -> str:
        """Повертає replacement для одного placeholder без зміни поточного DOCX contract."""
        nonlocal underline_needed
        key = match.group(0)
        value = params.get(key, None)

        if key == "{{CONDUTTORE_CF}}":
            # Для цього поля current шаблон покладається саме на "сирі" підкреслення,
            # а не на вирівнювання через пробіли, тому special-case лишається окремим.
            underline_needed = True
            length = underscored.get(key, 0)
            if value:
                return str(value)
            return "_" * max(length, 0)

        if key in underscored:
            underline_needed = True
            length = underscored[key]
            return fill_underscored(value, length)

        if value is not None:
            return str(value)

        return ""

    new_text = PLACEHOLDER_RE.sub(repl, text)
    if new_text != text:
        run.text = new_text
        if underline_needed:
            run.font.underline = WD_UNDERLINE.SINGLE


def replace_in_paragraph(paragraph, params: dict, underscored: dict):
    """Обробляє всі run-и в параграфі."""
    for run in paragraph.runs:
        _replace_in_run(run, params, underscored)


def replace_in_cell(cell, params: dict, underscored: dict):
    """Обробляє всі параграфи в комірці таблиці."""
    for paragraph in cell.paragraphs:
        replace_in_paragraph(paragraph, params, underscored)


def fill_attestazione_template(
    template_path: str,
    output_folder: str,
    filename: str,
    params: dict,
    underscored: dict,
) -> str:
    """
    Копіює DOCX-шаблон, підставляє поточні placeholder-значення й зберігає результат.

    Current document-generation semantics і output shape у migration slice не змінюються.
    """
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    out_path = output_folder / filename
    shutil.copy(template_path, out_path)

    doc = Document(out_path)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, params, underscored)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, params, underscored)

    doc.save(out_path)
    return str(out_path)


underscored = {
    "{{LOCATORE_NOME}}": 40,
    "{{LOCATORE_CF}}": 25,
    "{{LOCATORE_COMUNE_RES}}": 27,
    "{{LOCATORE_VIA}}": 27,
    "{{LOCATORE_CIVICO}}": 4,
    "{{IMMOBILE_COMUNE}}": 24,
    "{{IMMOBILE_VIA}}": 27,
    "{{IMMOBILE_CIVICO}}": 4,
    "{{IMMOBILE_PIANO}}": 4,
    "{{IMMOBILE_INTERNO}}": 4,
    "{{CONTRATTO_DATA}}": 13,
    "{{CONDUTTORE_NOME}}": 27,
    "{{CONDUTTORE_CF}}": 21,
    "{{CONDUTTORE_COMUNE}}": 24,
    "{{CONDUTTORE_VIA}}": 27,
    "{{DECORRENZA_DATA}}": 18,
    "{{REGISTRAZIONE_DATA}}": 13,
    "{{REGISTRAZIONE_NUM}}": 4,
    "{{AGENZIA_ENTRATE_SEDE}}": 25,
}


__all__ = [
    "fill_attestazione_template",
    "fill_underscored",
    "replace_in_paragraph",
    "replace_in_cell",
    "underscored",
]
